#!/usr/bin/env python3
"""Generate editable DOCX forms for locale copy review.

The Markdown review documents remain the canonical extracted source. This
script turns each locale file into reviewer-friendly tables while preserving
exact KEY/SOURCE anchors so corrected DOCX uploads can be parsed later.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


LOCALES = {
    "cs-CZ": "Czech",
    "en-EU": "English (EU)",
    "it-IT": "Italian",
}


@dataclass
class Entry:
    section: str
    subsection: str
    anchor_type: str
    anchor: str
    context: str
    flags: str
    text: str


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_paragraph(
    paragraph,
    size: int = 8,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05


def set_cell_text(
    cell,
    text: str,
    size: int = 8,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        style_paragraph(paragraph, size=size, bold=bold, color=color)


def parse_markdown(path: Path) -> list[Entry]:
    lines = path.read_text(encoding="utf-8").splitlines()
    entries: list[Entry] = []
    section = ""
    subsection = ""
    current: dict[str, str] | None = None
    text_lines: list[str] = []
    in_text = False

    def finish() -> None:
        nonlocal current, text_lines, in_text
        if not current:
            return
        anchor_type = current.get("anchor_type")
        anchor = current.get("anchor")
        if anchor_type and anchor:
            entries.append(
                Entry(
                    section=section,
                    subsection=subsection,
                    anchor_type=anchor_type,
                    anchor=anchor,
                    context=current.get("context", ""),
                    flags=current.get("flags", ""),
                    text="\n".join(text_lines).strip(),
                )
            )
        current = None
        text_lines = []
        in_text = False

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("## "):
            finish()
            section = line[3:]
            subsection = ""
            continue
        if line.startswith("### "):
            finish()
            subsection = line[4:]
            continue
        if line == "---":
            finish()
            continue
        if line.startswith("KEY: "):
            finish()
            current = {"anchor_type": "KEY", "anchor": line[5:]}
            continue
        if line.startswith("SOURCE: "):
            finish()
            current = {"anchor_type": "SOURCE", "anchor": line[8:]}
            continue
        if current is None:
            continue
        if line.startswith("CONTEXT: "):
            current["context"] = line[9:]
            in_text = False
            continue
        if line.startswith("FLAGS: "):
            current["flags"] = line[7:]
            in_text = False
            continue
        if line == "TEXT:":
            in_text = True
            text_lines = []
            continue
        if in_text:
            text_lines.append(line)
    finish()
    return entries


def grouped_by_section(entries: list[Entry]) -> dict[str, list[Entry]]:
    grouped: dict[str, list[Entry]] = {}
    for entry in entries:
        grouped.setdefault(entry.section, []).append(entry)
    return grouped


def add_table_for_section(doc: Document, entries: list[Entry]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    headers = [
        ("Reference", 1.35),
        ("Context / flags", 2.15),
        ("Current text", 3.15),
        ("Correct?", 0.9),
        ("Correction / query", 3.15),
    ]
    for index, (label, width) in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, width)
        set_cell_text(cell, label, size=8, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, "0F766E")

    for entry in entries:
        row = table.add_row()
        anchor_text = f"{entry.anchor_type}: {entry.anchor}"
        if entry.subsection:
            anchor_text = f"{anchor_text}\nGroup: {entry.subsection}"
        set_cell_text(row.cells[0], anchor_text, size=7)
        set_cell_text(row.cells[1], "\n".join(x for x in [entry.context, entry.flags] if x), size=7)
        set_cell_text(row.cells[2], entry.text, size=8)
        set_cell_text(row.cells[3], "☐ OK", size=8, bold=True)
        set_cell_text(
            row.cells[4],
            "If not OK, write:\nCORRECTION: ...\nor\nQUERY: ...",
            size=7,
            color=RGBColor(0x94, 0xA3, 0xB8),
        )

    doc.add_paragraph("")


def build_docx(locale: str, language: str, entries: list[Entry], output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8)
    normal.paragraph_format.space_after = Pt(2)

    title = doc.add_paragraph()
    run = title.add_run(f"{language} Text Review — splnit.eu")
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "Reviewer form. Keep the Reference column unchanged. Use the Correct? column to mark OK, or write CORRECTION:/QUERY: in the final column."
    )
    style_paragraph(subtitle, size=9, color=RGBColor(0x47, 0x55, 0x69))

    instructions = doc.add_table(rows=2, cols=3)
    instructions.style = "Table Grid"
    instruction_rows = [
        ("If correct", "Correct? column", "Replace ☐ OK with OK or ☑ OK."),
        ("If not correct", "Correction / query column", "Write CORRECTION: followed by the exact replacement, or QUERY: followed by the question."),
    ]
    for row_index, values in enumerate(instruction_rows):
        for col_index, value in enumerate(values):
            cell = instructions.rows[row_index].cells[col_index]
            set_cell_width(cell, [1.4, 1.8, 7.2][col_index])
            set_cell_text(cell, value, size=8, bold=col_index == 0)
            if col_index == 0:
                set_cell_shading(cell, "ECFDF5")

    doc.add_paragraph("")

    for section_name, section_entries in grouped_by_section(entries).items():
        heading = doc.add_paragraph()
        run = heading.add_run(section_name)
        run.font.name = "Arial"
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(4)
        add_table_for_section(doc, section_entries)

    for doc_section in doc.sections:
        footer = doc_section.footer.paragraphs[0]
        footer.text = f"{language} copy review form. Preserve Reference values. Mark OK or provide CORRECTION:/QUERY:."
        style_paragraph(footer, size=7, color=RGBColor(0x94, 0xA3, 0xB8))
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.core_properties.title = f"{language} Text Review — splnit.eu"
    doc.core_properties.subject = f"Editable table-form review document for {locale} web app copy"
    doc.core_properties.author = "Splnit.eu"
    doc.core_properties.comments = f"Generated from docs/review/text-review-{locale}.md; preserve Reference anchors for automated correction processing."
    doc.save(output)


def main() -> None:
    for locale, language in LOCALES.items():
        source = Path(f"docs/review/text-review-{locale}.md")
        output = Path(f"docs/review/text-review-{locale}.docx")
        entries = parse_markdown(source)
        if not entries:
            raise SystemExit(f"No entries parsed from {source}")
        build_docx(locale, language, entries, output)
        print(f"Wrote {output} with {len(entries)} review rows")


if __name__ == "__main__":
    main()
